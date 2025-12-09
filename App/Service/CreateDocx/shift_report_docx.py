from docx import Document
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any, Tuple
from io import BytesIO
from datetime import datetime
from pathlib import Path


def set_cell_text(cell, text, bold=False, align=None, font_size=12.6):
    """
    تعيين نص الخلية مع تنسيق بسيط
    """
    if not getattr(cell, "paragraphs", None):
        cell.add_paragraph()
    
    p = cell.paragraphs[0]
    # مسح المحتوى القديم
    for run in list(p.runs):
        r = run._r
        p._p.remove(r)
    
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    
    if align:
        if align == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'left':
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif align == 'right':
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def create_shift_report_in_memory(report_data: Dict[str, Any]) -> Tuple[BytesIO, str]:
    """
    إنشاء تقرير تقفيل الشفت DOCX في الذاكرة (بدون حفظ على القرص)
    
    Args:
        report_data: بيانات التقرير
    
    Returns:
        Tuple[BytesIO, str]: (ملف التقرير في الذاكرة, اسم الملف)
    """
    
    # إنشاء مستند جديد
    doc = Document()
    
    # ============================
    # ضبط مقاسات الطابعة الحرارية (نفس مقاسات فواتير الطلبات)
    # ============================
    section = doc.sections[0]
    section.page_width = Mm(80)      # عرض الورق الحراري
    section.page_height = Mm(297)      # طول الورق
    section.left_margin = Mm(4)        # هامش ضيق
    section.right_margin = Mm(4)
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)
    
    # ضبط الخط الافتراضي
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12.6)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0
    
    # استخراج البيانات
    shift_info = report_data['shift_info']
    orders_stats = report_data['orders_stats']
    financial_stats = report_data['financial_stats']
    payment_methods = report_data['payment_methods']
    
    # ============================
    # 0. اللوجو (إذا كان موجوداً)
    # ============================
    logo_path = Path(__file__).parent.parent.parent / "Static_Data" / "logo.png"
    if logo_path.exists():
        logo_paragraph = doc.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_paragraph.paragraph_format.space_after = Pt(5)
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(str(logo_path), width=Cm(3.0))  # عرض 3 سم (مناسب للطابعة الحرارية)
    
    # ============================
    # 1. العنوان الرئيسي
    # ============================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("تقرير تقفيل الشفت")
    title_run.font.size = Pt(16.8)
    title_run.font.bold = True
    
    # خط فاصل
    separator = doc.add_paragraph()
    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
    separator.paragraph_format.space_before = Pt(2)
    separator.paragraph_format.space_after = Pt(3)
    sep_run = separator.add_run("=" * 30)
    sep_run.font.size = Pt(11.2)
    
    # ============================
    # 2. معلومات الشفت الأساسية
    # ============================
    shift_header = doc.add_paragraph()
    shift_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    shift_header.paragraph_format.space_before = Pt(2)
    shift_header.paragraph_format.space_after = Pt(2)
    shift_header_run = shift_header.add_run("📋 معلومات الشفت")
    shift_header_run.font.size = Pt(14)
    shift_header_run.font.bold = True
    
    # جدول معلومات الشفت
    shift_table = doc.add_table(rows=5, cols=2)
    shift_table.style = 'Table Grid'
    
    # تنسيق وقت البداية والنهاية
    start_time_str = shift_info['start_time'].strftime("%I:%M %p") if shift_info['start_time'] else "---"
    end_time_str = shift_info['end_time'].strftime("%I:%M %p") if shift_info['end_time'] else "مفتوح"
    duration_str = f"{shift_info['duration_hours']} ساعة"
    
    set_cell_text(shift_table.cell(0, 0), shift_info['shift_number'], font_size=12.6)
    set_cell_text(shift_table.cell(0, 1), "رقم الشفت", bold=True, font_size=12.6)
    
    set_cell_text(shift_table.cell(1, 0), shift_info['shift_date'].strftime("%Y-%m-%d"), font_size=12.6)
    set_cell_text(shift_table.cell(1, 1), "التاريخ", bold=True, font_size=12.6)
    
    set_cell_text(shift_table.cell(2, 0), start_time_str, font_size=12.6)
    set_cell_text(shift_table.cell(2, 1), "وقت البداية", bold=True, font_size=12.6)
    
    set_cell_text(shift_table.cell(3, 0), end_time_str, font_size=12.6)
    set_cell_text(shift_table.cell(3, 1), "وقت النهاية", bold=True, font_size=12.6)
    
    set_cell_text(shift_table.cell(4, 0), duration_str, font_size=12.6)
    set_cell_text(shift_table.cell(4, 1), "مدة الشفت", bold=True, font_size=12.6)
    
    # ============================
    # 3. إحصائيات الطلبات
    # ============================
    orders_header = doc.add_paragraph()
    orders_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    orders_header.paragraph_format.space_before = Pt(3)
    orders_header.paragraph_format.space_after = Pt(2)
    orders_header_run = orders_header.add_run("📊 إحصائيات الطلبات")
    orders_header_run.font.size = Pt(14)
    orders_header_run.font.bold = True
    
    # جدول إحصائيات الطلبات
    orders_table = doc.add_table(rows=3, cols=2)
    orders_table.style = 'Table Grid'
    
    set_cell_text(orders_table.cell(0, 0), str(orders_stats['total_orders']), font_size=12.6)
    set_cell_text(orders_table.cell(0, 1), "إجمالي الطلبات", bold=True, font_size=12.6)
    
    set_cell_text(orders_table.cell(1, 0), str(orders_stats['delivered_orders']), font_size=12.6)
    set_cell_text(orders_table.cell(1, 1), "الطلبات المكتملة", bold=True, font_size=12.6)
    
    set_cell_text(orders_table.cell(2, 0), str(orders_stats['cancelled_orders']), font_size=12.6)
    set_cell_text(orders_table.cell(2, 1), "الطلبات الملغاة", bold=True, font_size=12.6)
    
    # ============================
    # 4. الملخص المالي
    # ============================
    financial_header = doc.add_paragraph()
    financial_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    financial_header.paragraph_format.space_before = Pt(3)
    financial_header.paragraph_format.space_after = Pt(2)
    financial_header_run = financial_header.add_run("💰 الملخص المالي")
    financial_header_run.font.size = Pt(14)
    financial_header_run.font.bold = True
    
    # جدول الملخص المالي
    financial_table = doc.add_table(rows=4, cols=2)
    financial_table.style = 'Table Grid'
    
    set_cell_text(financial_table.cell(0, 0), f"{financial_stats['total_sales']:.2f} ج.م", font_size=12.6)
    set_cell_text(financial_table.cell(0, 1), "إجمالي المبيعات", bold=True, font_size=12.6)
    
    set_cell_text(financial_table.cell(1, 0), f"{financial_stats['total_delivery_fees']:.2f} ج.م", font_size=12.6)
    set_cell_text(financial_table.cell(1, 1), "رسوم التوصيل", bold=True, font_size=12.6)
    
    set_cell_text(financial_table.cell(2, 0), f"{financial_stats['products_value']:.2f} ج.م", font_size=12.6)
    set_cell_text(financial_table.cell(2, 1), "قيمة المنتجات", bold=True, font_size=12.6)
    
    set_cell_text(financial_table.cell(3, 0), f"{financial_stats['average_order_value']:.2f} ج.م", font_size=12.6)
    set_cell_text(financial_table.cell(3, 1), "متوسط قيمة الطلب", bold=True, font_size=12.6)
    
    # ============================
    # 5. توزيع طرق الدفع
    # ============================
    if payment_methods and len(payment_methods) > 0:
        payment_header = doc.add_paragraph()
        payment_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        payment_header.paragraph_format.space_before = Pt(3)
        payment_header.paragraph_format.space_after = Pt(2)
        payment_header_run = payment_header.add_run("💳 توزيع طرق الدفع")
        payment_header_run.font.size = Pt(14)
        payment_header_run.font.bold = True
        
        # جدول توزيع طرق الدفع (4 أعمدة)
        payment_table = doc.add_table(rows=len(payment_methods) + 1, cols=4)
        payment_table.style = 'Table Grid'
        payment_table.autofit = False
        payment_table.allow_autofit = False
        
        # تحديد عرض الأعمدة
        payment_table.columns[0].width = Cm(1.5)  # النسبة
        payment_table.columns[1].width = Cm(1.8)  # المبلغ
        payment_table.columns[2].width = Cm(1.2)  # العدد
        payment_table.columns[3].width = Cm(2.0)  # الطريقة
        
        # رأس الجدول
        hdr_cells = payment_table.rows[0].cells
        set_cell_text(hdr_cells[0], "النسبة", bold=True, align='center', font_size=12.6)
        set_cell_text(hdr_cells[1], "المبلغ", bold=True, align='center', font_size=12.6)
        set_cell_text(hdr_cells[2], "العدد", bold=True, align='center', font_size=12.6)
        set_cell_text(hdr_cells[3], "الطريقة", bold=True, align='center', font_size=12.6)
        
        # بيانات طرق الدفع
        for idx, pm in enumerate(payment_methods, start=1):
            row_cells = payment_table.rows[idx].cells
            
            set_cell_text(row_cells[0], f"{pm['percentage']:.1f}%", font_size=12.6)
            set_cell_text(row_cells[1], f"{pm['total_amount']:.2f} ج.م", font_size=12.6)
            set_cell_text(row_cells[2], str(pm['orders_count']), font_size=12.6)
            set_cell_text(row_cells[3], pm['payment_method'], font_size=12.6)
    
    # ============================
    # 6. الخاتمة
    # ============================
    separator2 = doc.add_paragraph()
    separator2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    separator2.paragraph_format.space_before = Pt(3)
    separator2.paragraph_format.space_after = Pt(2)
    sep2_run = separator2.add_run("=" * 30)
    sep2_run.font.size = Pt(11.2)
    
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer_run = footer.add_run(f"تم إنشاء التقرير: {datetime.now().strftime('%Y-%m-%d %I:%M %p')}")
    footer_run.font.size = Pt(11.2)
    
    # ============================
    # 7. إنشاء اسم الملف وحفظه في الذاكرة
    # ============================
    date_str = shift_info['shift_date'].strftime("%Y-%m-%d")
    filename = f"Shift-Report-{shift_info['shift_number']}-{date_str}.docx"
    
    # حفظ في الذاكرة (BytesIO)
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)  # العودة لبداية الملف للقراءة
    
    return file_stream, filename
