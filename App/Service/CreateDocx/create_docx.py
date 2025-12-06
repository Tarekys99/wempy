from docx import Document
from docx.shared import Pt, Mm, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any, Tuple
from io import BytesIO
from pathlib import Path
import os


def set_cell_text(cell, text, bold=False, align=None, font_size=9):
    """
    تعيين نص الخلية مع تنسيق بسيط
    """
    if not getattr(cell, "paragraphs", None):
        cell.add_paragraph()
    
    p = cell.paragraphs[0]
    # مسح المحتوى القديم
    for run in list(p.runs):
        r = run._r
        p._p.remove(r)
    
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    
    if align:
        if align == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'left':
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif align == 'right':
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def create_invoice_in_memory(invoice_data: Dict[str, Any]) -> Tuple[BytesIO, str]:
    """
    إنشاء فاتورة DOCX في الذاكرة (بدون حفظ على القرص)
    
    Args:
        invoice_data: بيانات الفاتورة
    
    Returns:
        Tuple[BytesIO, str]: (ملف الفاتورة في الذاكرة, اسم الملف)
    """
    
    # إنشاء مستند جديد
    doc = Document()
    
    # ============================
    # ضبط مقاسات الطابعة الحرارية
    # ============================
    section = doc.sections[0]
    section.page_width = Mm(72.1)      # عرض الورق الحراري
    section.page_height = Mm(297)      # طول الورق
    section.left_margin = Mm(4)        # هامش ضيق
    section.right_margin = Mm(4)
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)
    
    # ضبط الخط الافتراضي
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0
    
    # ============================
    # 0. اللوجو (إذا كان موجوداً)
    # ============================
    logo_path = Path(__file__).parent.parent.parent / "Static_Data" / "logo.png"
    if logo_path.exists():
        logo_paragraph = doc.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_paragraph.paragraph_format.space_after = Pt(5)
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(str(logo_path), width=Cm(3.0))  # عرض 3.5 سم (مناسب للطابعة الحرارية)
    
    # ============================
    # 1. العنوان الرئيسي (رقم الطلب)
    # ============================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"Order {invoice_data['order_number']}")
    title_run.font.size = Pt(12)
    title_run.font.bold = True
    
    # التاريخ
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_before = Pt(0)
    date_p.paragraph_format.space_after = Pt(2)
    date_run = date_p.add_run(f"{invoice_data['order_date']}")
    date_run.font.size = Pt(8)
    
    # خط فاصل
    separator = doc.add_paragraph()
    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
    separator.paragraph_format.space_before = Pt(0)
    separator.paragraph_format.space_after = Pt(2)
    sep_run = separator.add_run("-" * 30)
    sep_run.font.size = Pt(8)
    
    # ============================
    # 2. معلومات الشفت
    # ============================
    shift_info = doc.add_paragraph()
    shift_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shift_info.paragraph_format.space_before = Pt(0)
    shift_info.paragraph_format.space_after = Pt(1)
    shift_run = shift_info.add_run(f"SHIFT - {invoice_data['shift_number']}")
    shift_run.font.size = Pt(9)
    shift_run.font.bold = True
    
    # ============================
    # 3. معلومات المستلم
    # ============================
    separator2 = doc.add_paragraph()
    separator2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    separator2.paragraph_format.space_before = Pt(0)
    separator2.paragraph_format.space_after = Pt(3)
    sep2_run = separator2.add_run("-" * 30)
    sep2_run.font.size = Pt(8)
    
    recipient_header = doc.add_paragraph()
    recipient_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    recipient_header.paragraph_format.space_before = Pt(0)
    recipient_header.paragraph_format.space_after = Pt(2)
    recipient_header_run = recipient_header.add_run("بيانات العميل")
    recipient_header_run.font.size = Pt(10)
    recipient_header_run.font.bold = True
    
    # جدول بيانات العميل
    customer_table = doc.add_table(rows=4, cols=2)
    customer_table.style = 'Table Grid'
    
    set_cell_text(customer_table.cell(0, 0), invoice_data['recipient_name'], font_size=9)
    set_cell_text(customer_table.cell(0, 1), "الاسم", bold=True, font_size=9)
    set_cell_text(customer_table.cell(1, 0), invoice_data['recipient_phone'], font_size=9)
    set_cell_text(customer_table.cell(1, 1), "الهاتف", bold=True, font_size=9)
    
    address_text = f"{invoice_data['city']}، {invoice_data['street']}، {invoice_data['building']}"
    set_cell_text(customer_table.cell(2, 0), address_text, font_size=8)
    set_cell_text(customer_table.cell(2, 1), "العنوان", bold=True, font_size=9)
    
    # إضافة اسم المنطقة
    set_cell_text(customer_table.cell(3, 0), invoice_data['zone_name'], font_size=9)
    set_cell_text(customer_table.cell(3, 1), "المنطقة", bold=True, font_size=9)
    
    # ============================
    # 4. جدول المنتجات
    # ============================
    products_header = doc.add_paragraph()
    products_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    products_header.paragraph_format.space_before = Pt(3)
    products_header.paragraph_format.space_after = Pt(2)
    products_header_run = products_header.add_run("تفاصيل الطلب")
    products_header_run.font.size = Pt(10)
    products_header_run.font.bold = True
    
    # إنشاء جدول المنتجات (4 أعمدة)
    items_count = len(invoice_data['items'])
    products_table = doc.add_table(rows=items_count + 1, cols=4)
    products_table.style = 'Table Grid'
    
    # تحديد عرض الأعمدة (المجموع = 6.0 سم ليتناسب مع باقي الجداول)
    col_widths = [Cm(1.2), Cm(1.0), Cm(0.6), Cm(3.2)]  # الإجمالي، السعر، الكمية، الصنف
    
    for row in products_table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width
    
    # رأس الجدول
    hdr_cells = products_table.rows[0].cells
    set_cell_text(hdr_cells[0], "الإجمالي", bold=True, align='center', font_size=9)
    set_cell_text(hdr_cells[1], "السعر", bold=True, align='center', font_size=9)
    set_cell_text(hdr_cells[2], "الكمية", bold=True, align='center', font_size=9)
    set_cell_text(hdr_cells[3], "الصنف", bold=True, align='center', font_size=9)
    
    # بيانات المنتجات
    for idx, item in enumerate(invoice_data['items'], start=1):
        row_cells = products_table.rows[idx].cells
        
        # دمج اسم المنتج مع المتغير (في سطرين منفصلين)
        product_full = f"{item['product_name']}"
        if item['variant_info']:
            product_full += f"\n({item['variant_info']})"
        
        set_cell_text(row_cells[0], f"{item['subtotal']:.2f}", font_size=9)
        set_cell_text(row_cells[1], f"{item['unit_price']:.2f}", font_size=9)
        set_cell_text(row_cells[2], str(item['quantity']), font_size=9)
        set_cell_text(row_cells[3], product_full, font_size=8)
    
    # ============================
    # 5. ملاحظات الطلب
    # ============================
    if invoice_data['order_notes'] and invoice_data['order_notes'] != "لا توجد ملاحظات":
        notes_header = doc.add_paragraph()
        notes_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        notes_header.paragraph_format.space_before = Pt(3)
        notes_header.paragraph_format.space_after = Pt(2)
        notes_header_run = notes_header.add_run("ملاحظات العميل")
        notes_header_run.font.size = Pt(9)
        notes_header_run.font.bold = True
        
        notes_table = doc.add_table(rows=1, cols=1)
        notes_table.style = 'Table Grid'
        notes_table.autofit = False
        notes_table.allow_autofit = False
        notes_table.columns[0].width = Cm(6.0)
        set_cell_text(notes_table.cell(0, 0), invoice_data['order_notes'], align='right', font_size=8)
    
    # ============================
    # 5.1 الملاحظات الخارجية
    # ============================
    if invoice_data.get('external_notes') and invoice_data['external_notes']:
        external_notes_header = doc.add_paragraph()
        external_notes_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        external_notes_header.paragraph_format.space_before = Pt(3)
        external_notes_header.paragraph_format.space_after = Pt(2)
        external_notes_header_run = external_notes_header.add_run("ملاحظات خارجية")
        external_notes_header_run.font.size = Pt(9)
        external_notes_header_run.font.bold = True
        
        external_notes_table = doc.add_table(rows=1, cols=1)
        external_notes_table.style = 'Table Grid'
        external_notes_table.autofit = False
        external_notes_table.allow_autofit = False
        external_notes_table.columns[0].width = Cm(6.0)
        set_cell_text(external_notes_table.cell(0, 0), invoice_data['external_notes'], align='right', font_size=8)
    
    # ============================
    # 6. الحساب النهائي
    # ============================
    payment_header = doc.add_paragraph()
    payment_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    payment_header.paragraph_format.space_before = Pt(3)
    payment_header.paragraph_format.space_after = Pt(2)
    payment_header_run = payment_header.add_run("الحساب")
    payment_header_run.font.size = Pt(10)
    payment_header_run.font.bold = True
    
    # جدول الملخص المالي
    # تحديد عدد الصفوف بناءً على طريقة الدفع
    rows_count = 5 if invoice_data.get('payment_id') == 2 else 4
    summary_table = doc.add_table(rows=rows_count, cols=2)
    summary_table.style = 'Table Grid'
    
    # مجموع المنتجات
    set_cell_text(summary_table.cell(0, 0), f"{invoice_data['items_subtotal']:.2f} ج.م", font_size=9)
    set_cell_text(summary_table.cell(0, 1), "المجموع الفرعي", bold=True, font_size=9)
    
    # تكلفة التوصيل
    set_cell_text(summary_table.cell(1, 0), f"{invoice_data['delivery_fee']:.2f} ج.م", font_size=9)
    set_cell_text(summary_table.cell(1, 1), "رسوم التوصيل", bold=True, font_size=9)
    
    # الإجمالي النهائي
    set_cell_text(summary_table.cell(2, 0), f"{invoice_data['total_price']:.2f} ج.م", bold=True, font_size=10)
    set_cell_text(summary_table.cell(2, 1), "الإجمالي النهائي", bold=True, font_size=10)
    
    # طريقة الدفع
    set_cell_text(summary_table.cell(3, 0), invoice_data['payment_method'], font_size=9)
    set_cell_text(summary_table.cell(3, 1), "طريقة الدفع", bold=True, font_size=9)
    
    # رقم التحويل (فقط إذا كانت طريقة الدفع محفظة إلكترونية)
    if invoice_data.get('payment_id') == 2:
        set_cell_text(summary_table.cell(4, 0), "01008403545", font_size=9)
        set_cell_text(summary_table.cell(4, 1), "رقم التحويل فودافون كاش", bold=True, font_size=9)
    
    # ============================
    # 7. الخاتمة
    # ============================
    separator3 = doc.add_paragraph()
    separator3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    separator3.paragraph_format.space_before = Pt(3)
    separator3.paragraph_format.space_after = Pt(2)
    sep3_run = separator3.add_run("-" * 30)
    sep3_run.font.size = Pt(8)
    
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer_run = footer.add_run("شكراً لتعاملكم معنا")
    footer_run = footer.add_run("\n٢٩٨ ن شارع العشرين - البوابة الرابعة - حدائق الاهرام")
    footer_run.font.size = Pt(9)
    footer_run.font.bold = True
    
    # ============================
    # 8. إنشاء اسم الملف وحفظه في الذاكرة
    # ============================
    date_str = (
        invoice_data['order_date']
        .replace('/', '-')
        .replace(' ', '_')
        .replace(':', '-')
    )
    filename = f"ORDER-{invoice_data['order_number']}-SHIFT-{invoice_data['shift_number']}.docx"
    
    # حفظ في الذاكرة (BytesIO)
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)  # العودة لبداية الملف للقراءة
    
    return file_stream, filename